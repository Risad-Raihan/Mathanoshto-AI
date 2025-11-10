"""
File parsing and text extraction utilities
"""
from pathlib import Path
from typing import Optional, Dict, Any
import json
import xml.etree.ElementTree as ET
from datetime import datetime


class FileParser:
    """Parse and extract text from various file types"""
    
    @staticmethod
    def parse_pdf(file_path: Path) -> Dict[str, Any]:
        """Extract text from PDF using pdfplumber (better than PyPDF2)"""
        try:
            import pdfplumber
            
            text_content = []
            metadata = {}
            
            with pdfplumber.open(file_path) as pdf:
                # Extract metadata
                if pdf.metadata:
                    metadata = {
                        'author': pdf.metadata.get('Author'),
                        'creator': pdf.metadata.get('Creator'),
                        'producer': pdf.metadata.get('Producer'),
                        'subject': pdf.metadata.get('Subject'),
                        'title': pdf.metadata.get('Title'),
                        'creation_date': pdf.metadata.get('CreationDate'),
                        'page_count': len(pdf.pages)
                    }
                
                # Extract text from all pages
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_content.append(text)
            
            full_text = "\n\n".join(text_content)
            
            return {
                'success': True,
                'text': full_text,
                'metadata': metadata,
                'page_count': metadata.get('page_count', 0)
            }
        
        except Exception as e:
            return {
                'success': False,
                'error': str(e),
                'text': '',
                'metadata': {}
            }
    
    @staticmethod
    def parse_docx(file_path: Path) -> Dict[str, Any]:
        """Extract text from DOCX files"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            
            # Extract text from paragraphs
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            full_text = "\n".join(text_content)
            
            # Extract metadata
            metadata = {}
            if hasattr(doc, 'core_properties'):
                props = doc.core_properties
                metadata = {
                    'author': props.author,
                    'title': props.title,
                    'subject': props.subject,
                    'created': str(props.created) if props.created else None,
                    'modified': str(props.modified) if props.modified else None,
                }
            
            return {
                'success': True,
                'text': full_text,
                'metadata': metadata
            }
        
        except Exception as e:
            return {
                'success': False,
                'error': str(e),
                'text': '',
                'metadata': {}
            }
    
    @staticmethod
    def parse_txt(file_path: Path) -> Dict[str, Any]:
        """Extract text from plain text files"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        text = f.read()
                    
                    return {
                        'success': True,
                        'text': text,
                        'metadata': {'encoding': encoding}
                    }
                except UnicodeDecodeError:
                    continue
            
            return {
                'success': False,
                'error': 'Could not decode file with supported encodings',
                'text': '',
                'metadata': {}
            }
        
        except Exception as e:
            return {
                'success': False,
                'error': str(e),
                'text': '',
                'metadata': {}
            }
    
    @staticmethod
    def parse_csv(file_path: Path) -> Dict[str, Any]:
        """Parse CSV files and extract data"""
        try:
            import pandas as pd
            
            # Read CSV
            df = pd.read_csv(file_path)
            
            # Convert to text representation
            text_content = []
            
            # Add column headers
            text_content.append("Columns: " + ", ".join(df.columns.tolist()))
            text_content.append("")
            
            # Add row count
            text_content.append(f"Total rows: {len(df)}")
            text_content.append("")
            
            # Add first few rows as text
            text_content.append("Sample data:")
            text_content.append(df.head(10).to_string())
            
            # Add basic statistics for numeric columns
            numeric_cols = df.select_dtypes(include=['number']).columns
            if len(numeric_cols) > 0:
                text_content.append("")
                text_content.append("Numeric column statistics:")
                text_content.append(df[numeric_cols].describe().to_string())
            
            full_text = "\n".join(text_content)
            
            metadata = {
                'rows': len(df),
                'columns': len(df.columns),
                'column_names': df.columns.tolist(),
                'dtypes': {col: str(dtype) for col, dtype in df.dtypes.items()}
            }
            
            return {
                'success': True,
                'text': full_text,
                'metadata': metadata,
                'dataframe': df  # Keep for potential later use
            }
        
        except Exception as e:
            return {
                'success': False,
                'error': str(e),
                'text': '',
                'metadata': {}
            }
    
    @staticmethod
    def parse_json(file_path: Path) -> Dict[str, Any]:
        """Parse JSON files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # Convert JSON to readable text
            text = json.dumps(data, indent=2, ensure_ascii=False)
            
            metadata = {
                'type': type(data).__name__,
                'size': len(str(data))
            }
            
            if isinstance(data, dict):
                metadata['keys'] = list(data.keys())
            elif isinstance(data, list):
                metadata['items'] = len(data)
            
            return {
                'success': True,
                'text': text,
                'metadata': metadata,
                'data': data
            }
        
        except Exception as e:
            return {
                'success': False,
                'error': str(e),
                'text': '',
                'metadata': {}
            }
    
    @staticmethod
    def parse_xml(file_path: Path) -> Dict[str, Any]:
        """Parse XML files"""
        try:
            tree = ET.parse(file_path)
            root = tree.getroot()
            
            # Convert XML to text representation
            def element_to_text(element, level=0):
                indent = "  " * level
                text_parts = [f"{indent}<{element.tag}>"]
                
                if element.text and element.text.strip():
                    text_parts.append(f"{indent}  {element.text.strip()}")
                
                for child in element:
                    text_parts.append(element_to_text(child, level + 1))
                
                text_parts.append(f"{indent}</{element.tag}>")
                return "\n".join(text_parts)
            
            text = element_to_text(root)
            
            metadata = {
                'root_tag': root.tag,
                'attributes': root.attrib
            }
            
            return {
                'success': True,
                'text': text,
                'metadata': metadata
            }
        
        except Exception as e:
            return {
                'success': False,
                'error': str(e),
                'text': '',
                'metadata': {}
            }
    
    @staticmethod
    def parse_image_ocr(file_path: Path) -> Dict[str, Any]:
        """Extract text from images using OCR"""
        try:
            import pytesseract
            from PIL import Image
            
            # Open image
            image = Image.open(file_path)
            
            # Extract text using OCR
            text = pytesseract.image_to_string(image)
            
            # Get image metadata
            metadata = {
                'format': image.format,
                'mode': image.mode,
                'size': image.size,
                'width': image.width,
                'height': image.height
            }
            
            # Get EXIF data if available
            if hasattr(image, '_getexif') and image._getexif():
                exif = image._getexif()
                if exif:
                    metadata['exif'] = {k: str(v) for k, v in exif.items()}
            
            return {
                'success': True,
                'text': text,
                'metadata': metadata
            }
        
        except Exception as e:
            return {
                'success': False,
                'error': str(e),
                'text': '',
                'metadata': {}
            }
    
    @staticmethod
    def get_image_metadata(file_path: Path) -> Dict[str, Any]:
        """Get image metadata without OCR"""
        try:
            from PIL import Image
            
            image = Image.open(file_path)
            
            metadata = {
                'format': image.format,
                'mode': image.mode,
                'size': image.size,
                'width': image.width,
                'height': image.height
            }
            
            return {
                'success': True,
                'metadata': metadata
            }
        
        except Exception as e:
            return {
                'success': False,
                'error': str(e),
                'metadata': {}
            }
    
    @staticmethod
    def parse_file(file_path: Path, file_type: str, enable_ocr: bool = False) -> Dict[str, Any]:
        """
        Parse file based on type
        
        Args:
            file_path: Path to the file
            file_type: Type of file ('pdf', 'docx', 'txt', etc.)
            enable_ocr: Whether to enable OCR for images
        
        Returns:
            Dictionary with 'success', 'text', 'metadata', and optionally 'error'
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            return {
                'success': False,
                'error': 'File not found',
                'text': '',
                'metadata': {}
            }
        
        # Route to appropriate parser
        if file_type == 'pdf':
            return FileParser.parse_pdf(file_path)
        elif file_type == 'docx':
            return FileParser.parse_docx(file_path)
        elif file_type == 'txt':
            return FileParser.parse_txt(file_path)
        elif file_type == 'csv':
            return FileParser.parse_csv(file_path)
        elif file_type == 'json':
            return FileParser.parse_json(file_path)
        elif file_type == 'xml':
            return FileParser.parse_xml(file_path)
        elif file_type == 'image':
            if enable_ocr:
                return FileParser.parse_image_ocr(file_path)
            else:
                return FileParser.get_image_metadata(file_path)
        else:
            return {
                'success': False,
                'error': f'Unsupported file type: {file_type}',
                'text': '',
                'metadata': {}
            }


# Convenience instance
file_parser = FileParser()

